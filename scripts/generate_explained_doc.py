#!/usr/bin/env python3
"""Generate balanced BCA+MA Invoicing Module Explained doc — offline, no DB push."""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
TMP = Path("/tmp")
OUT = ROOT / "docs" / "Invoicing_Module_Explained.docx"

PRIMARY = RGBColor(0x0F, 0x2A, 0x44)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)  # sky for balanced doc (distinct from proposal amber)
ACCENT_HEX = "0EA5E9"
GREY_DARK = RGBColor(0x1E, 0x29, 0x3B)
GREY = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x94, 0xA3, 0xB8)
BORDER = "CBD5E1"
ALT_ROW = "F1F5F9"
HEADER_BG = "0F2A44"
TODAY = date.today().strftime("%d %B %Y")

def set_cell_shading(cell, hex):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), hex); shd.set(qn("w:val"), "clear"); cell._tc.get_or_add_tcPr().append(shd)
def set_cell_border(cell, **kw):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tb = tcPr.first_child_found_in("w:tcBorders")
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tcPr.append(tb)
    for e in ("top","left","bottom","right","insideH","insideV"):
        d=kw.get(e)
        if d:
            el=tb.find(qn(f"w:{e}"))
            if el is None:
                el=OxmlElement(f"w:{e}"); tb.append(el)
            for k in ("val","sz","space","color"):
                if k in d: el.set(qn(f"w:{k}"), str(d[k]))
def add_hline(p, color=ACCENT_HEX):
    pr=p._p.get_or_add_pPr(); bdr=OxmlElement("w:pBdr"); b=OxmlElement("w:bottom"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6"); b.set(qn("w:space"),"1"); b.set(qn("w:color"),color); bdr.append(b); pr.append(bdr)

def style_doc(doc):
    s=doc.styles["Normal"]; s.font.name="Calibri"; s.font.size=Pt(10); s.font.color.rgb=GREY_DARK; s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.15
    for name, sz, col, bold in [("Heading 1",18,PRIMARY,True),("Heading 2",13,PRIMARY,True),("Heading 3",11,GREY,True)]:
        h=doc.styles[name]; h.font.name="Calibri"; h.font.size=Pt(sz); h.font.color.rgb=col; h.font.bold=bold; h.paragraph_format.space_before=Pt(16 if "1" in name else 10); h.paragraph_format.space_after=Pt(5)
        if name=="Heading 3": h.font.italic=True

def add_hdr_ftr(doc):
    sec=doc.sections[0]; sec.top_margin=Cm(1.4); sec.bottom_margin=Cm(1.4); sec.left_margin=Cm(1.6); sec.right_margin=Cm(1.6); sec.header_distance=Cm(0.8); sec.footer_distance=Cm(0.8)
    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=hp.add_run("PROKON ERP  •  Invoicing Module — Balanced Guide (BCA + MA)  •  CONFIDENTIAL"); r.font.size=Pt(7); r.font.color.rgb=MUTED; r.font.italic=True
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=fp.add_run(); a=OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"),"begin"); b=OxmlElement("w:instrText"); b.set(qn("xml:space"),"preserve"); b.text="PAGE"; c=OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"),"end"); r._r.append(a); r2=fp.add_run(); r2._r.append(b); r3=fp.add_run(); r3._r.append(c); r=fp.add_run(f"  •  invoicing-module  •  {TODAY}"); r.font.size=Pt(7); r.font.color.rgb=MUTED

def cover(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=p.add_run("PROKON ERP  •  KNOWLEDGE BASE"); r.font.size=Pt(7.5); r.font.color.rgb=ACCENT; r.font.bold=True; p.paragraph_format.space_after=Pt(18)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run("  EXPLAINED  •  BALANCED  •  BCA + MA  "); r.font.size=Pt(7); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.bold=True; pr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),HEADER_BG); shd.set(qn("w:val"),"clear"); pr.append(shd); p.paragraph_format.space_after=Pt(14)
    p=doc.add_paragraph(); r=p.add_run("Invoicing Module\n— How the Bill Works\nFlow • Tables • Integration"); r.font.size=Pt(28); r.font.color.rgb=PRIMARY; r.bold=True; p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=0.96
    add_hline(doc.add_paragraph(), ACCENT_HEX)
    p=doc.add_paragraph(); r=p.add_run("One bill, many departments — written so a BCA student (how it’s built) and an MA student (why it matters) both get it."); r.font.size=Pt(9.5); r.font.color.rgb=GREY; r.italic=True; p.paragraph_format.space_after=Pt(14)
    tbl=doc.add_table(rows=1,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    c0,c1=tbl.cell(0,0),tbl.cell(0,1); c0.width=Inches(3.3); c1.width=Inches(3.3)
    for lab,val in [("Branch","invoicing-module"),("Date",TODAY),("For","BCA + MA  •  Finance + Operations"),("Level","Balanced — Simple first, Technical next"),("Docs","Flow  •  Tables  •  Integration  •  UI/Safety"),("Guarantee","Migrations add only — zero rows deleted")]:
        pp=c0.add_paragraph(); rr=pp.add_run(lab+"  "); rr.font.size=Pt(7); rr.font.color.rgb=MUTED; rr.bold=True; rr=pp.add_run(val); rr.font.size=Pt(8); rr.font.color.rgb=GREY_DARK; pp.paragraph_format.space_after=Pt(1)
    pp=c1.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=pp.add_run(" HOW TO READ "); rr.font.size=Pt(8); rr.bold=True; rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); pr=pp._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),ACCENT_HEX); shd.set(qn("w:val"),"clear"); pr.append(shd)
    for t in ["MA: read the blue boxes + first para of each chapter (story).","BCA: read the grey code lines + tables (files & SQL).","Both: the purple flow diagram is the shared map."]:
        pp=c1.add_paragraph(style="List Bullet"); rr=pp.add_run(t); rr.font.size=Pt(7.5); rr.font.color.rgb=GREY_DARK; pp.paragraph_format.space_after=Pt(1)
    for c in (c0,c1):
        set_cell_shading(c,"F8FAFC" if c==c0 else "FFF7ED"); set_cell_border(c,top={"val":"single","sz":"4","color":BORDER},left={"val":"single","sz":"4","color":BORDER},bottom={"val":"single","sz":"4","color":BORDER},right={"val":"single","sz":"4","color":BORDER})
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); r=p.add_run("Tip: "); r.bold=True; r.font.size=Pt(7.5); r.font.color.rgb=PRIMARY; r=p.add_run("Each chapter starts with a blue MA box (analogy), then a BCA table/diagram. If you only have 10 minutes, read the blue boxes + the flow on page 2."); r.font.size=Pt(7.5); r.font.color.rgb=GREY
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16); r=p.add_run("—  SIMPLE FIRST  •  TECHNICAL NEXT  —"); r.font.size=Pt(7); r.font.color.rgb=MUTED; r.italic=True
    doc.add_page_break()

def toc(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("CONTENTS"); r.font.size=Pt(11); r.bold=True; r.font.color.rgb=PRIMARY; p.paragraph_format.space_after=Pt(4); add_hline(doc.add_paragraph(), BORDER)
    for num,title,pg in [("1","Flow & Working — The Staged Bill (Balanced)","3"),("2","Table Structure — Registers Behind the Bill","6"),("3","Integrations — How the Bill Talks to Others","9"),("4","UI, Print & Safety — What You See & What Guards It","12"),("5","Glossary, FAQ & Appendix (Migrations Safe)","14")]:
        p=doc.add_paragraph(); pp=p._p.get_or_add_pPr(); tabs=OxmlElement("w:tabs"); tab=OxmlElement("w:tab"); tab.set(qn("w:val"),"right"); tab.set(qn("w:leader"),"dot"); tab.set(qn("w:pos"),"9350"); tabs.append(tab); pp.append(tabs); r=p.add_run(f"{num}   {title}"); r.font.size=Pt(9); r.font.color.rgb=PRIMARY; r.bold=True; r=p.add_run(f"\t{pg}"); r.font.size=Pt(8); r.font.color.rgb=MUTED
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8); r=p.add_run("Balanced reading path: MA → blue boxes + diagrams. BCA → tables + file:line code. Both → flow + glossary. Zero rows are deleted by migrations — they only ADD columns/tables."); r.font.size=Pt(7); r.font.color.rgb=MUTED; r.italic=True
    doc.add_page_break()

def heading(doc,text,lvl=1):
    h=doc.add_heading(text,level=lvl)
    for r in h.runs:
        if r is not None and hasattr(r, 'font'):
            r.font.name="Calibri"
    if lvl==1: add_hline(doc.add_paragraph(), ACCENT_HEX)
    return h

def styled_tbl(doc, headers, rows, widths=None, hdr=HEADER_BG):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False if widths else True
    if widths:
        for i,w in enumerate(widths): t.columns[i].width=Inches(w)
    hdr_cells=t.rows[0].cells
    for i,h in enumerate(headers):
        hdr_cells[i].text=""; p=hdr_cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(h); r.font.size=Pt(7); r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); hdr_cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_shading(hdr_cells[i],hdr); set_cell_border(hdr_cells[i],top={"val":"single","sz":"4","color":BORDER},left={"val":"single","sz":"4","color":BORDER},bottom={"val":"single","sz":"4","color":BORDER},right={"val":"single","sz":"4","color":BORDER})
    for idx,row in enumerate(rows):
        cs=t.add_row().cells
        for j,v in enumerate(row):
            cs[j].text=""; p=cs[j].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(str(v)); r.font.size=Pt(7.5); r.font.color.rgb=GREY_DARK; cs[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p.paragraph_format.space_after=Pt(1)
            if idx%2==1: set_cell_shading(cs[j],ALT_ROW)
            set_cell_border(cs[j],top={"val":"single","sz":"3","color":BORDER},left={"val":"single","sz":"3","color":BORDER},bottom={"val":"single","sz":"3","color":BORDER},right={"val":"single","sz":"3","color":BORDER})
    doc.add_paragraph().paragraph_format.space_after=Pt(4); return t

def ma_box(doc, title, text):
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; c=tbl.cell(0,0); set_cell_shading(c,"EFF6FF"); set_cell_border(c,top={"val":"single","sz":"6","color":ACCENT_HEX},left={"val":"single","sz":"6","color":ACCENT_HEX},bottom={"val":"single","sz":"6","color":ACCENT_HEX},right={"val":"single","sz":"6","color":ACCENT_HEX})
    p=c.paragraphs[0]; r=p.add_run("MA — " + title + "  "); r.bold=True; r.font.size=Pt(8); r.font.color.rgb=PRIMARY; r=p.add_run(text); r.font.size=Pt(9); r.font.color.rgb=GREY_DARK; p.paragraph_format.space_after=Pt(2); doc.add_paragraph().paragraph_format.space_after=Pt(2)

def bca_box(doc, title, text):
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; c=tbl.cell(0,0); set_cell_shading(c,"F8FAFC"); set_cell_border(c,top={"val":"single","sz":"4","color":BORDER},left={"val":"single","sz":"4","color":BORDER},bottom={"val":"single","sz":"4","color":BORDER},right={"val":"single","sz":"4","color":BORDER})
    p=c.paragraphs[0]; r=p.add_run("BCA — " + title + "  "); r.bold=True; r.font.size=Pt(7.5); r.font.color.rgb=GREY; r=p.add_run(text); r.font.size=Pt(7.5); r.font.color.rgb=GREY_DARK; p.paragraph_format.space_after=Pt(2); doc.add_paragraph().paragraph_format.space_after=Pt(2)

def code(doc, txt):
    p=doc.add_paragraph(); pr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),"F8FAFC"); shd.set(qn("w:val"),"clear"); pr.append(shd); bdr=OxmlElement("w:pBdr")
    for e in ("top","left","bottom","right"):
        el=OxmlElement(f"w:{e}"); el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4"); el.set(qn("w:space"),"4"); el.set(qn("w:color"),BORDER); bdr.append(el)
    pr.append(bdr); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    for line in txt.strip().split("\n"):
        pp=doc.add_paragraph(); pr2=pp._p.get_or_add_pPr(); shd2=OxmlElement("w:shd"); shd2.set(qn("w:fill"),"F8FAFC"); shd2.set(qn("w:val"),"clear"); pr2.append(shd2)
        pp.paragraph_format.space_after=Pt(0); pp.paragraph_format.space_before=Pt(0); pp.paragraph_format.left_indent=Inches(0.14)
        r=pp.add_run(line if line.strip() else " "); r.font.name="Consolas"; r.font.size=Pt(7); r.font.color.rgb=GREY_DARK
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def build():
    doc=Document(); style_doc(doc); add_hdr_ftr(doc); cover(doc); toc(doc)
    # Read markdown sections generated by subagents
    for path, title in [("/tmp/inv_section_flow.md","1. Flow & Working"),("/tmp/inv_section_tables.md","2. Table Structure"),("/tmp/inv_section_integration.md","3. Integrations"),("/tmp/inv_section_ui.md","4. UI, Print & Safety")]:
        # Heading duplicate guard: sections already contain heading, so just render their content as paragraphs/tables via simple markdown pass
        # For brevity we will add a top heading then include raw markdown as formatted paragraphs
        # Instead we will parse markdown lightly: headings -> heading, tables -> styled_tbl, code -> code
        # Quick: read file and add as preformatted + styled
        heading(doc, title, lvl=1)
        txt = Path(path).read_text(encoding="utf-8") if Path(path).exists() else f"(missing {path})"
        # Remove the first markdown heading line if present to avoid double (the section title we already added)
        lines = txt.splitlines()
        if lines and lines[0].strip().startswith("##"):
            lines = lines[1:]
        txt = "\n".join(lines).strip()
        # Parse line-by-line to preserve subheadings (###, ####) and code fences and tables
        # Split into blocks but keep heading lines separate
        raw_blocks = [b.strip() for b in txt.split("\n\n") if b.strip()]
        for b in raw_blocks:
            # Handle code fence block ```...```
            if b.strip().startswith("```"):
                # extract inner lines between fences
                inner = b.strip()
                # remove first line ``` and last ```
                if inner.startswith("```"):
                    inner = inner.split("\n",1)[1] if "\n" in inner else ""
                    if inner.rstrip().endswith("```"):
                        inner = inner.rsplit("```",1)[0]
                code(doc, inner.strip())
                continue
            # Split block into lines to detect subheadings inside
            sub_lines = b.splitlines()
            # If block is a single subheading line like "### 3.2 Upstream"
            if len(sub_lines)==1 and sub_lines[0].strip().startswith("###"):
                lvl = 3 if sub_lines[0].strip().startswith("### ") else 3
                heading(doc, sub_lines[0].strip().lstrip("#").strip(), lvl=lvl)
                continue
            # If block starts with ### then first line is heading, rest is paragraph
            if sub_lines and sub_lines[0].strip().startswith("###"):
                heading(doc, sub_lines[0].strip().lstrip("#").strip(), lvl=3)
                # remainder of block after heading
                remainder = "\n".join(sub_lines[1:]).strip()
                if not remainder:
                    continue
                b = remainder
            # Handle bold-only MA/BCA boxes: detect leading **MA** or **BCA**
            if b.strip().startswith("**") or b.lstrip().startswith("MA —") or b.lstrip().startswith("BCA —"):
                # keep as styled box if contains MA/BCA marker
                if "MA —" in b[:30] or b.strip().startswith("MA"):
                    ma_box(doc, "Simple", b.replace("**","").replace("MA —","").replace(">","").strip()[:1200])
                    continue
                if "BCA —" in b[:30] or b.strip().startswith("BCA"):
                    bca_box(doc, "Technical", b.replace("**","").replace("BCA —","").strip()[:1600])
                    continue
            # Markdown table block
            if b.strip().startswith("|") and "|" in b:
                rows = [r.strip() for r in b.splitlines() if r.strip().startswith("|")]
                if len(rows) >= 2:
                    hdr = [h.strip().replace("**","") for h in rows[0].strip("|").split("|")]
                    start = 2 if set(rows[1].replace("|","").replace("-","").replace(":","").strip()) == set() or "---" in rows[1] else 1
                    data = [[c.strip().replace("**","") for c in r.strip("|").split("|")] for r in rows[start:]]
                    clean = []
                    for r in data:
                        if not any(c for c in r): continue
                        if len(r) < len(hdr): r = r + [""]*(len(hdr)-len(r))
                        if len(r) > len(hdr): r = r[:len(hdr)]
                        clean.append(r)
                    data = clean
                    if hdr and data:
                        styled_tbl(doc, hdr, data)
                        continue
            # Bullet list block starting with - or *
            if b.strip().startswith("- ") or b.strip().startswith("* "):
                for line in b.splitlines():
                    if line.strip().startswith("- ") or line.strip().startswith("* "):
                        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2)
                        r = p.add_run(line.strip()[2:].replace("**","").strip())
                        r.font.size=Pt(9); r.font.color.rgb=GREY_DARK
                    else:
                        if line.strip():
                            p=doc.add_paragraph(line.strip().replace("**","").strip()); p.paragraph_format.space_after=Pt(3); 
                            for rr in p.runs: rr.font.size=Pt(9)
                continue
            # Default: handle inline bold ** and add as paragraph with mixed runs
            # Simple: split by ** and alternate bold
            p = doc.add_paragraph()
            p.paragraph_format.space_after=Pt(4)
            # crude bold handling: split on **
            parts = b.split("**")
            for i, part in enumerate(parts):
                if not part: continue
                # odd index means inside ** (bold)
                r = p.add_run(part.strip())
                r.font.size=Pt(9); r.font.color.rgb=GREY_DARK
                if i % 2 == 1:
                    r.bold = True; r.font.color.rgb=PRIMARY
            # If no **, ensure size
            if "**" not in b:
                for r in p.runs: r.font.size=Pt(9)
        # add page break except last
        if "Safety" not in title:
            doc.add_page_break()

    # 5 Glossary & Appendix
    heading(doc, "5. Glossary, FAQ & Appendix (Safe Migrations)", lvl=1)
    ma_box(doc, "Glossary — MA", "B2B = business to business (buyer has GSTIN). B2C = buyer is a person (no GSTIN, shown as URP). HSN = product code, GSTIN = seller/buyer tax ID, IRN = 64-hex bill fingerprint from govt, LUT = letter for SEZ export without tax, POS = place of supply state, FY = April-March year, URP = unregistered person.")
    bca_box(doc, "FAQ — BCA", "Q: Does migration delete rows? A: No — all migrations are IF NOT EXISTS / DEFAULT backfill / CREATE OR REPLACE (zero DELETE/TRUNCATE, append-only invoice_print_log). Q: Why JSONB for transport_details? A: 25-field Tally parity stored as one GIN-indexed object (vehicle/distance/PIN) — easy to query via transport_details->>'vehicle_no' with index. Q: How is numbering gap-free? A: SELECT ... FOR UPDATE + pg_advisory_xact_lock per branch+FY serializes next_seq. Q: Why 50000 threshold inclusive? A: E-Way required when total ≥50000 per spec.")
    styled_tbl(doc, ["Migrations you run (order matters)","What each adds (safe)","Rows deleted?"], [
        ["20260902000000_invoicing_staged.sql","sales_type enum 7, transport_details JSONB, einvoice/eway status, compliance_json, print_log table, triggers","0 — DEFAULTs fill existing 7 invoices"],
        ["20260902000001_fix_h10_supply_class_guard.sql","nil↔local_nil_rated / zero_rated↔sez_zero_rated trigger","0"],
        ["20260902000002_harden_invoicing_rls.sql","RLS sales.* policies for settings/items/eway","0 — policy DDL only"],
        ["20260902000003_gstin_lut_hardening.sql","GSTIN checksum + LUT trigger","0"],
        ["20260902000004_print_audit_rpc.sql","increment_invoice_print() RPC atomic","0"],
        ["then: supabase gen types","regens src/integrations/supabase/types.ts","0"],
    ], widths=[2.0,2.2,0.9])
    ma_box(doc, "Try it — 5-minute lab", "Create invoice: Pick Branch (GSTIN) + Customer (GSTIN) → Sales Type Local-ItemWise → Transport Self, Vehicle HR26AA1234, Gurugram 122001, 120km → Items 2× Batteries → Save → Generate GST JSON → download → paste mock IRN (64-hex) → Generated → Print Original/Duplicate/Triplicate. Head Sales KPI shows IRN pending → Complete.")
    bca_box(doc, "Files changed", "supabase/migrations/20260902000000..04 (5 files), src/lib/{sales,gst,india,transport,invoiceJson,einvoice,invoicePdf}, src/components/{TransportDetailsModal,InvoicePrintModal}, src/routes/_app/sales.{invoices.new,$id,index,settings}, scripts/seed_invoicing_demo.ts — all on invoicing-module, branch not pushed.")
    # Signoff
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(12); r=p.add_run("—  End — Balanced Guide — "); r.font.size=Pt(8); r.bold=True; r.font.color.rgb=PRIMARY
    styled_tbl(doc, ["Role","Name","Date"], [["Prokon — Finance/CA","___________________________","____ / ____ / 2026"],["Prokon — Operations","___________________________","____ / ____ / 2026"],["Engineering — BCA Owner","___________________________","____ / ____ / 2026"]], widths=[1.5,2.0,1.5])
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6); r=p.add_run("Generated python-docx on invoicing-module • No DB push — you run migrations manually • zero rows deleted"); r.font.size=Pt(6); r.font.color.rgb=MUTED
    OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(str(OUT)); return OUT

if __name__ == "__main__":
    out = build()
    print(f"Wrote {out} ({out.stat().st_size} bytes)")
    from docx import Document as D2
    d=D2(str(out))
    print(f"Verify: {len(d.paragraphs)} paras, {len(d.tables)} tables, headings={[p.text for p in d.paragraphs if p.style.name.startswith('Heading')][:10]}")
