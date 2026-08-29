#!/usr/bin/env python3
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x25, 0x63, 0xEB)
PRIMARY_HEX = "2563EB"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_HEX = "F1F5FD"
MUTED = RGBColor(0x47, 0x55, 0x69)

CATEGORIES = [
    ("Finance & Billing", [
        ("B-01", "A delivery challan that errors mid-conversion still stays 'unbilled', so it can be invoiced a second time - the customer gets two invoices for one delivery.",
         "Challan to Invoice conversion logic", "Mark challan 'Converted' only after the invoice and its items are safely saved; surface the intermediate error instead of swallowing it."),
        ("B-02", "Double-clicking 'Convert' (quote to order to challan to invoice) silently creates duplicate sales orders/invoices that look normal.",
         "Document conversion actions", "Disable the convert button during save and add a 'done' marker to every conversion step so a retry cannot recreate it."),
        ("B-03", "The header discount is never split into saved line items, so the invoice taxable total and the item-wise sum disagree - GST filings do not reconcile.",
         "Invoice header / line-item tax calc", "Apportion the header discount into each line item (paise-exact) before saving taxable values; GST charged post-discount."),
        ("B-04", "Quotations decide tax type from the state name while invoices use the GSTIN code, so a quote shows CGST+SGST but the invoice charges IGST for the same deal.",
         "Quote & Invoice GST engines", "Use one intra/inter-state rule via GSTIN (isIntraSupply) in both quote and invoice engines."),
        ("B-05", "A payment can be allocated across more dues than it covers, and concurrent entries can apply the same money twice.",
         "Payment allocation service", "Cap each allocation to the invoice's remaining balance and revalidate inside a single server transaction."),
        ("B-06", "Overselling is only blocked on the screen; any authenticated user could POST an invoice with allow_negative_stock and sell stock that does not exist.",
         "supabase/migrations/20260823100000_bugfix_hardening.sql", "DB trigger assert_negative_stock_admin() enforces allow_negative_stock as admin-only at the database level (invoices + both challan tables)."),
    ]),
    ("Stock & Inventory", [
        ("B-07", "Before issuing a serial to a ticket/invoice the system does not re-check it is still free, so two warranties can point at one physical machine.",
         "issueStockToTicket / serial issue flow", "Claim the serial first and refuse already-issued serials with a guarded update before assigning."),
        ("B-08", "A goods-received note already marked 'Submitted' (already added to stock) can be reopened and re-saved, inflating stock counts.",
         "GRN form + DB policies", "Block GRN edits once status is Submitted; allow changes only via the admin reverse flow (client + DB)."),
        ("B-09", "A blank or garbage GRN quantity quietly becomes 1 unit instead of stopping you - you receive 0 but stock goes up by 1.",
         "GRN quantity validation", "Treat empty/NaN quantity as invalid and block saving until corrected."),
        ("B-10", "A fractional quantity on a serial-tracked line slips through (e.g. 2.5 units with 2 serials), so records and physical units mismatch.",
         "Invoice line-item validation", "Require whole-number quantity on serialized lines and match the serial count exactly."),
        ("B-11", "List pages stop fetching past ~1,000 rows with no error - challans/GRNs/IMS records exist in the DB but silently disappear from view.",
         "Challan / GRN / IMS list pages", "Use the existing paged fetch (or server-side pagination) for large lists."),
    ]),
    ("Dates & Time (IST)", [
        ("B-12", "Documents created between midnight and 5:30 AM get yesterday's date because 'today' is computed in UTC, not Indian time.",
         "Date helpers (istTodayIso / localDateIso)", "Replace UTC-based dates with IST helpers across ~30 call sites so invoices/POs/payments/indents use Asia/Kolkata."),
        ("B-13", "AMC contracts ending 'today' can show as already expired (or still active) depending on the UTC hour.",
         "AMC expiry logic", "Compare contract dates using IST strings instead of UTC midnight."),
        ("B-14", "Returnable challans are flagged 'overdue' about 5.5 hours early, triggering premature follow-ups.",
         "Returnable overdue check", "Evaluate overdue using IST date terms, not a UTC timestamp."),
        ("B-15", "Ticket SLA/response timers use the device timezone, so field staff abroad and HQ see different 'truths'.",
         "Ticket SLA computation", "Pin all SLA computations to IST."),
    ]),
    ("Silent Failures", [
        ("B-16", "Several steps hide their failures (invoice/GDC forms, writers, ticket log, deletes) so users see 'Success!' while nothing was saved.",
         "Invoice/GDC forms, writers, ticket log, deletes", "Remove silent catches; surface the error to the user and retry or roll back instead of swallowing it."),
        ("B-17", "Fire-and-forget actions (logout tracking, deletes) can fail with no message, leaving gaps in audit trails.",
         "Logout / delete handlers", "Await these calls and report failure instead of ignoring it."),
    ]),
    ("Admin & Security", [
        ("B-18", "Admin status was sometimes just a browser label, so sensitive actions (payroll corrections, GRN delete, stock reverse) could be reachable by non-admins.",
         "RLS policies + SECURITY DEFINER RPCs", "Enforce admin/role checks in database policies; payroll/GRN-delete/challan-delete paths already admin-gated (residual verified)."),
        ("B-19", "The first person to click 'Claim Admin' could self-promote without real verification; the check could race.",
         "supabase/migrations/20260823100000_bugfix_hardening.sql", "guard_first_admin_claim() takes an advisory xact lock and re-checks before allowing the bootstrap admin claim."),
    ]),
    ("Races & Concurrency", [
        ("B-20", "Approving a stock transfer does not check its current status, so a double-click or two admins can approve it twice.",
         "Transfer approval flow", "Update only if the current status matches the expected value (conditional update)."),
        ("B-21", "A delayed autosave in indents can land after the record was already changed again, clobbering the latest work.",
         "Indent autosave", "Guard in-flight saves with a version/abort token so stale writes cancel."),
        ("B-22", "Printing labels twice can fail to increment the printed count, so you can't tell what was actually printed.",
         "Print-count updater", "Use an atomic 'add 1' update with error handling."),
    ]),
    ("Payroll", [
        ("B-23", "Salary/advance math is complex and was never automatically tested, so a logic slip can silently pay the wrong amount to everyone.",
         "Payroll salary math + paired writes", "Added automated tests for salary math; wrapped paired writes (attendance+audit, payment+advance) in transactions with compensating rollback."),
        ("B-24", "Undoing attendance deletes records and updates an audit flag in two steps; if the second fails, payroll changes lose their audit trail.",
         "Attendance correction flow", "Ensure the audit update completes with, or rolls back alongside, the delete."),
    ]),
    ("CRM / Document Conversion", [
        ("B-25", "Only the first conversion step (Quote to SO) was protected against duplicates; SO to Challan and to Invoice were not, desyncing statuses.",
         "Conversion chain", "Add a 'converted' marker to every conversion step, not just the first."),
    ]),
    ("Print / PDF", [
        ("B-26", "The print/PDF view recomputes GST totals instead of reading stored values, so what the customer receives can differ from what's on file.",
         "Invoice print / PDF view", "Print from the stored breakdown values, not by recalculating."),
    ]),
    ("Document Header & Company Profile (final fixes)", [
        ("B-27", "Selected branch no longer exists - resolveHeader threw on a missing branch/warehouse instead of falling back to the company profile.",
         "src/lib/documentHeader.ts", "Make resolveHeader fall back to the company profile when the selected branch/warehouse no longer exists instead of throwing."),
        ("B-28", "Company Master incomplete: Company Master record not found - fetchCompanyProfile threw on a DB error, leaving the AMC company state null.",
         "src/lib/companyProfile.ts", "Harden fetchCompanyProfile to handle DB errors gracefully so the AMC company state is always populated, not left null."),
    ]),
]

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color

def add_heading_bar(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = PRIMARY
    run.font.name = "Calibri"
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), PRIMARY_HEX)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("Prokon ERP - Bug Fix Report")
    trun.bold = True
    trun.font.size = Pt(22)
    trun.font.color.rgb = PRIMARY
    trun.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("Hardening of Finance, Stock, Time, Security & Document flows")
    srun.font.size = Pt(11)
    srun.font.color.rgb = MUTED
    srun.italic = True

    total_bugs = sum(len(bugs) for _, bugs in CATEGORIES)
    today = date.today().strftime("%d %B %Y")

    summary = doc.add_table(rows=1, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_left = summary.rows[0].cells[0]
    s_right = summary.rows[0].cells[1]
    set_cell_text(s_left, "Total Bugs Documented:  " + str(total_bugs), bold=True, color=PRIMARY, size=11)
    set_cell_text(s_right, "Report Date:  " + today, bold=True, color=PRIMARY, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    shade_cell(s_left, ALT_ROW_HEX)
    shade_cell(s_right, ALT_ROW_HEX)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    headers = ["#", "Bug Description", "File(s) Affected", "Fix Applied"]
    widths = [Inches(0.5), Inches(3.0), Inches(1.9), Inches(2.1)]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=HEADER_TEXT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[i], PRIMARY_HEX)
        hdr[i].width = widths[i]

    row_idx = 0
    for cat_name, bugs in CATEGORIES:
        add_heading_bar(doc, cat_name)
        for bid, desc, files, fix in bugs:
            row = table.add_row().cells
            set_cell_text(row[0], bid, bold=True, color=PRIMARY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[1], desc, size=9)
            set_cell_text(row[2], files, size=8, color=MUTED)
            set_cell_text(row[3], fix, size=9)
            for i in range(4):
                row[i].width = widths[i]
            if row_idx % 2 == 1:
                for i in range(4):
                    shade_cell(row[i], ALT_ROW_HEX)
            row_idx += 1

    doc.add_paragraph()
    note = doc.add_paragraph()
    nrun = note.add_run(
        "Source: BUG_REPORT.md (B-01-B-26, code-review leads) and final hardening commits "
        "(resolveHeader / fetchCompanyProfile). Items B-01-B-26 were addressed in the "
        "money/stock/time hardening pass; B-06, B-18, B-19 and related sweep items were "
        "enforced at the database level via supabase migration 20260823100000_bugfix_hardening.sql."
    )
    nrun.font.size = Pt(8)
    nrun.font.color.rgb = MUTED
    nrun.italic = True

    out_path = "/Users/jai/Desktop/Prokon Erp/Bug_Fix_Report.docx"
    doc.save(out_path)
    print("Saved: " + out_path)
    print("Total bugs documented: " + str(total_bugs))

if __name__ == "__main__":
    main()
